import io
import os
import shutil
import tempfile
import zipfile

import fitz  # PyMuPDF
try:
    import ghostscript
except RuntimeError:
    ghostscript = None
import openpyxl
from flask import (
    Flask,
    jsonify,
    render_template_string,
    request,
    send_file,
)
from pdf2docx import Converter
from pdf2docx.converter import ConversionException
import pytesseract
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100MB max
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["OUTPUT_FOLDER"] = "outputs"

# Criar diretórios se não existirem
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "xlsx", "jpg", "jpeg", "png"}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# Template HTML
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>LocalPDF - Academic PDF Utility</title>
    <style>
        :root {
            /* Palette: Clean Academic/Product */
            --bg-body: #f8f9fa;
            --bg-surface: #ffffff;
            --bg-surface-hover: #f1f3f5;
            --text-main: #212529;
            --text-muted: #495057;
            --border-color: #dee2e6;
            --accent-primary: #339af0;
            --accent-primary-hover: #228be6;
            --accent-success: #40c057;
            --accent-success-hover: #2b8a3e;
            --accent-danger: #fa5252;
            
            --space-xs: 4px;
            --space-sm: 8px;
            --space-md: 16px;
            --space-lg: 24px;
            --space-xl: 32px;
            --space-2xl: 48px;
            --space-3xl: 64px;

            --radius-sm: 4px;
            --radius-md: 6px;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
            background-color: var(--bg-body);
            color: var(--text-main);
            line-height: 1.5;
            min-height: 100vh;
            display: flex;
            flex-direction: column;
        }

        .navbar {
            background-color: var(--bg-surface);
            border-bottom: 1px solid var(--border-color);
            padding: var(--space-md) var(--space-xl);
            display: flex;
            align-items: center;
            justify-content: space-between;
        }

        .brand {
            font-size: 1.25rem;
            font-weight: 600;
            letter-spacing: -0.01em;
            color: var(--text-main);
            text-decoration: none;
            display: flex;
            align-items: center;
            gap: var(--space-sm);
        }

        .brand svg {
            width: 20px;
            height: 20px;
            color: var(--accent-primary);
        }

        .container {
            max-width: 1000px;
            margin: 0 auto;
            padding: var(--space-2xl) var(--space-lg);
            flex: 1;
            width: 100%;
        }

        .header {
            margin-bottom: var(--space-2xl);
        }

        .header h1 {
            font-size: 2rem;
            font-weight: 600;
            letter-spacing: -0.02em;
            margin-bottom: var(--space-xs);
        }

        .header p {
            color: var(--text-muted);
            font-size: 1.05rem;
        }

        .tools-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(280px, 1fr));
            gap: var(--space-md);
        }

        .tool-card {
            background: var(--bg-surface);
            border: 1px solid var(--border-color);
            border-radius: var(--radius-md);
            padding: var(--space-lg);
            cursor: pointer;
            transition: all 0.15s ease;
            display: flex;
            flex-direction: column;
            gap: var(--space-sm);
        }

        .tool-card:hover {
            border-color: var(--accent-primary);
            background-color: var(--bg-surface-hover);
        }

        .tool-header {
            display: flex;
            align-items: center;
            gap: var(--space-sm);
        }

        .tool-icon {
            color: var(--accent-primary);
            width: 20px;
            height: 20px;
            display: flex;
            align-items: center;
            justify-content: center;
        }

        .tool-icon svg {
            width: 100%;
            height: 100%;
        }

        .tool-card h3 {
            font-size: 1rem;
            font-weight: 600;
            color: var(--text-main);
        }

        .tool-card p {
            font-size: 0.875rem;
            color: var(--text-muted);
        }

        /* Tool View Setup */
        .hidden {
            display: none !important;
        }

        .view-header {
            display: flex;
            align-items: center;
            gap: var(--space-md);
            margin-bottom: var(--space-xl);
            padding-bottom: var(--space-md);
            border-bottom: 1px solid var(--border-color);
        }

        .back-btn {
            background: transparent;
            color: var(--text-muted);
            border: 1px solid var(--border-color);
            border-radius: var(--radius-sm);
            padding: var(--space-xs) var(--space-sm);
            cursor: pointer;
            font-size: 0.875rem;
            display: flex;
            align-items: center;
            gap: var(--space-xs);
            transition: all 0.15s ease;
        }

        .back-btn:hover {
            background: var(--bg-surface-hover);
            color: var(--text-main);
        }

        .workspace {
            background: var(--bg-surface);
            border: 1px solid var(--border-color);
            border-radius: var(--radius-md);
            padding: var(--space-xl);
        }

        .workspace-header {
            margin-bottom: var(--space-lg);
        }

        .workspace-header h2 {
            font-size: 1.25rem;
            font-weight: 600;
            margin-bottom: var(--space-xs);
        }

        .workspace-header p {
            color: var(--text-muted);
            font-size: 0.875rem;
        }

        .upload-area {
            border: 1px dashed var(--border-color);
            border-radius: var(--radius-md);
            padding: var(--space-3xl) var(--space-xl);
            text-align: center;
            background: var(--bg-body);
            transition: all 0.15s ease;
            cursor: pointer;
            margin-bottom: var(--space-lg);
        }

        .upload-area:hover, .upload-area.dragover {
            border-color: var(--accent-primary);
            background: rgba(51, 154, 240, 0.05);
        }

        .upload-icon {
            color: var(--text-muted);
            margin-bottom: var(--space-sm);
        }

        .file-input {
            display: none;
        }

        .btn-primary {
            background: var(--accent-primary);
            color: white;
            border: none;
            border-radius: var(--radius-sm);
            padding: var(--space-sm) var(--space-lg);
            font-size: 0.875rem;
            font-weight: 500;
            cursor: pointer;
            transition: background 0.15s ease;
        }

        .btn-primary:hover {
            background: var(--accent-primary-hover);
        }

        .btn-primary:disabled {
            background: var(--border-color);
            cursor: not-allowed;
            color: var(--text-muted);
        }

        .btn-success {
            background: var(--accent-success);
        }
        .btn-success:hover {
            background: var(--accent-success-hover);
        }

        .file-list {
            margin-bottom: var(--space-lg);
            display: flex;
            flex-direction: column;
            gap: var(--space-sm);
        }

        .file-item {
            display: flex;
            align-items: center;
            justify-content: space-between;
            background: var(--bg-body);
            border: 1px solid var(--border-color);
            padding: var(--space-sm) var(--space-md);
            border-radius: var(--radius-sm);
            font-size: 0.875rem;
        }

        .file-info {
            display: flex;
            align-items: center;
            gap: var(--space-sm);
        }

        .btn-remove {
            background: transparent;
            color: var(--accent-danger);
            border: none;
            cursor: pointer;
            font-size: 0.875rem;
        }
        .btn-remove:hover {
            text-decoration: underline;
        }

        .progress {
            width: 100%;
            background: var(--border-color);
            border-radius: var(--radius-sm);
            height: 4px;
            margin: var(--space-md) 0;
            overflow: hidden;
        }

        .progress-bar {
            height: 100%;
            background: var(--accent-primary);
            width: 0%;
            transition: width 0.3s ease;
        }

        .alert {
            padding: var(--space-md);
            border-radius: var(--radius-sm);
            margin-top: var(--space-md);
            font-size: 0.875rem;
        }

        .alert-success {
            background: rgba(64, 192, 87, 0.1);
            color: var(--accent-success-hover);
            border: 1px solid rgba(64, 192, 87, 0.2);
        }

        .alert-error {
            background: rgba(250, 82, 82, 0.1);
            color: var(--accent-danger);
            border: 1px solid rgba(250, 82, 82, 0.2);
        }

        .footer {
            border-top: 1px solid var(--border-color);
            padding: var(--space-lg) var(--space-xl);
            text-align: center;
            color: var(--text-muted);
            font-size: 0.75rem;
            background: var(--bg-surface);
        }
        
        .footer a {
            color: var(--text-muted);
            text-decoration: none;
            margin: 0 var(--space-xs);
        }
        
        .footer a:hover {
            color: var(--text-main);
        }

        .action-row {
            display: flex;
            justify-content: flex-end;
            gap: var(--space-sm);
            margin-top: var(--space-md);
            border-top: 1px solid var(--border-color);
            padding-top: var(--space-md);
        }
    </style>
</head>
<body>
    <nav class="navbar">
        <a href="#" class="brand" onclick="showHome()">
            <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"></path><polyline points="14 2 14 8 20 8"></polyline><line x1="16" y1="13" x2="8" y2="13"></line><line x1="16" y1="17" x2="8" y2="17"></line><polyline points="10 9 9 9 8 9"></polyline></svg>
            LocalPDF
        </a>
    </nav>

    <div class="container">
        <!-- Home View -->
        <div id="home-view">
            <div class="header">
                <h1>Local PDF Tools</h1>
                <p>Fast, offline document manipulation. No cloud uploads required.</p>
            </div>

            <div class="tools-grid" id="tools-grid-container">
                <!-- Injected via JS -->
            </div>
        </div>

        <!-- Tool View -->
        <div id="tool-views" class="hidden">
            <div class="view-header">
                <button class="back-btn" onclick="showHome()">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><line x1="19" y1="12" x2="5" y2="12"></line><polyline points="12 19 5 12 12 5"></polyline></svg>
                    Back to tools
                </button>
            </div>

            <div class="workspace">
                <div class="workspace-header">
                    <h2 id="tool-title">Tool Title</h2>
                    <p id="tool-description">Description of what the tool does.</p>
                </div>

                <div class="upload-area" id="upload-area" onclick="document.getElementById('file-input').click()">
                    <input type="file" id="file-input" class="file-input" multiple>
                    <div class="upload-icon">
                        <svg width="32" height="32" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"></path><polyline points="17 8 12 3 7 8"></polyline><line x1="12" y1="3" x2="12" y2="15"></line></svg>
                    </div>
                    <p>Click to select files or drag them here</p>
                    <p style="font-size: 0.75rem; margin-top: 4px; color: var(--text-muted);" id="tool-accept-text"></p>
                </div>

                <div id="file-list" class="file-list"></div>

                <div class="action-row">
                    <button id="convert-btn" class="btn-primary btn-success hidden" onclick="convertFiles()">Process Files</button>
                </div>

                <div id="progress" class="progress hidden">
                    <div id="progress-bar" class="progress-bar"></div>
                </div>
                <div style="text-align: center; font-size: 0.75rem; color: var(--text-muted);" id="progress-text" class="hidden">Processing locally...</div>

                <div id="result" class="hidden"></div>
            </div>
        </div>
    </div>

    <footer class="footer">
        <div>LocalPDF Utility — Fast, secure, and private.</div>
        <div style="margin-top: var(--space-xs);">
            Originally developed by Virgilio Borges | 
            <a href="https://github.com/virgiliojr94" target="_blank">GitHub</a> | 
            <a href="https://www.linkedin.com/in/virgiliojunior94/" target="_blank">LinkedIn</a> |
            <a href="mailto:virgilio.junior94@gmail.com">Contact</a>
        </div>
    </footer>

    <script>
        const ICONS = {
            image: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="3" width="18" height="18" rx="2" ry="2"></rect><circle cx="8.5" cy="8.5" r="1.5"></circle><polyline points="21 15 16 10 5 21"></polyline></svg>',
            file: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M13 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V9z"></path><polyline points="13 2 13 9 20 9"></polyline></svg>',
            merge: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M10 13a5 5 0 0 0 7.54.54l3-3a5 5 0 0 0-7.07-7.07l-1.72 1.71"></path><path d="M14 11a5 5 0 0 0-7.54-.54l-3 3a5 5 0 0 0 7.07 7.07l1.71-1.71"></path></svg>',
            split: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="6" cy="6" r="3"></circle><circle cx="6" cy="18" r="3"></circle><line x1="20" y1="4" x2="8.12" y2="15.88"></line><line x1="14.47" y1="14.48" x2="20" y2="20"></line><line x1="8.12" y1="8.12" x2="12" y2="12"></line></svg>',
            compress: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M21 16V8a2 2 0 0 0-1-1.73l-7-4a2 2 0 0 0-2 0l-7 4A2 2 0 0 0 3 8v8a2 2 0 0 0 1 1.73l7 4a2 2 0 0 0 2 0l7-4A2 2 0 0 0 21 16z"></path><polyline points="3.27 6.96 12 12.01 20.73 6.96"></polyline><line x1="12" y1="22.08" x2="12" y2="12"></line></svg>',
            lock: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="11" width="18" height="11" rx="2" ry="2"></rect><path d="M7 11V7a5 5 0 0 1 10 0v4"></path></svg>',
            text: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><line x1="17" y1="10" x2="3" y2="10"></line><line x1="21" y1="6" x2="3" y2="6"></line><line x1="21" y1="14" x2="3" y2="14"></line><line x1="17" y1="18" x2="3" y2="18"></line></svg>',
            refresh: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="23 4 23 10 17 10"></polyline><polyline points="1 20 1 14 7 14"></polyline><path d="M3.51 9a9 9 0 0 1 14.85-3.36L23 10M1 14l4.64 4.36A9 9 0 0 0 20.49 15"></path></svg>',
            search: '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="11" cy="11" r="8"></circle><line x1="21" y1="21" x2="16.65" y2="16.65"></line></svg>'
        };

        const tools = {
            'pdf-to-images': {
                title: 'PDF to Images',
                description: 'Convert each page of your PDF into separate images',
                accept: '.pdf',
                multiple: false,
                icon: 'image'
            },
            'images-to-pdf': {
                title: 'Images to PDF',
                description: 'Combine multiple images into a single PDF file',
                accept: '.jpg,.jpeg,.png',
                multiple: true,
                icon: 'file'
            },
            'merge-pdf': {
                title: 'Merge PDFs',
                description: 'Combine multiple PDF files into a single document',
                accept: '.pdf',
                multiple: true,
                icon: 'merge'
            },
            'split-pdf': {
                title: 'Split PDF',
                description: 'Extract specific pages from your PDF',
                accept: '.pdf',
                multiple: false,
                icon: 'split'
            },
            'compress-pdf': {
                title: 'Compress PDF',
                description: 'Reduce the PDF file size while maintaining quality',
                accept: '.pdf',
                multiple: false,
                icon: 'compress'
            },
            'pdf-to-pdfa': {
                title: 'PDF to PDF/A',
                description: 'Convert PDFs to the PDF/A-1b archiving standard',
                accept: '.pdf',
                multiple: true,
                icon: 'lock'
            },
            'word-to-pdf': {
                title: 'Word to PDF',
                description: 'Convert Word documents (.docx) to PDF',
                accept: '.docx',
                multiple: true,
                icon: 'file'
            },
            'excel-to-pdf': {
                title: 'Excel to PDF',
                description: 'Convert Excel spreadsheets (.xlsx) to PDF',
                accept: '.xlsx',
                multiple: false,
                icon: 'file'
            },
            'txt-to-pdf': {
                title: 'TXT to PDF',
                description: 'Convert plain text files (.txt) to PDF',
                accept: '.txt',
                multiple: false,
                icon: 'text'
            },
            'pdf-to-word': {
                title: 'PDF to Word',
                description: 'Convert your PDF documents to editable Word (.docx)',
                accept: '.pdf',
                multiple: false,
                icon: 'refresh'
            },
            'ocr-pdf': {
                title: 'OCR on PDF',
                description: 'Extract text from PDFs and scanned images using OCR',
                accept: '.pdf,.jpg,.jpeg,.png',
                multiple: false,
                icon: 'search'
            }
        };

        // Render Home Grid
        const grid = document.getElementById('tools-grid-container');
        for (const [id, tool] of Object.entries(tools)) {
            const el = document.createElement('div');
            el.className = 'tool-card';
            el.onclick = () => showTool(id);
            el.innerHTML = `
                <div class="tool-header">
                    <div class="tool-icon">${ICONS[tool.icon]}</div>
                    <h3>${tool.title}</h3>
                </div>
                <p>${tool.description}</p>
            `;
            grid.appendChild(el);
        }

        let currentTool = '';
        let uploadedFiles = [];

        function showTool(toolName) {
            currentTool = toolName;
            const tool = tools[toolName];

            document.getElementById('home-view').classList.add('hidden');
            document.getElementById('tool-views').classList.remove('hidden');
            document.getElementById('tool-title').innerText = tool.title;
            document.getElementById('tool-description').innerText = tool.description;
            document.getElementById('file-input').accept = tool.accept;
            document.getElementById('file-input').multiple = tool.multiple;
            
            let formats = tool.accept.split(',').map(ext => ext.toUpperCase().replace('.', '')).join(', ');
            document.getElementById('tool-accept-text').innerText = `Accepted formats: ${formats}`;

            uploadedFiles = [];
            updateFileList();
            hideResult();
        }

        function showHome() {
            document.getElementById('home-view').classList.remove('hidden');
            document.getElementById('tool-views').classList.add('hidden');
            uploadedFiles = [];
        }

        function updateFileList() {
            const fileList = document.getElementById('file-list');
            const convertBtn = document.getElementById('convert-btn');

            if (uploadedFiles.length === 0) {
                fileList.innerHTML = '';
                convertBtn.classList.add('hidden');
                return;
            }

            fileList.innerHTML = uploadedFiles.map((file, index) => `
                <div class="file-item">
                    <div class="file-info">
                        <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M13 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V9z"></path><polyline points="13 2 13 9 20 9"></polyline></svg>
                        <span>${file.name} <span style="color: var(--text-muted)">(${(file.size / 1024 / 1024).toFixed(2)} MB)</span></span>
                    </div>
                    <button class="btn-remove" onclick="removeFile(${index})">Remove</button>
                </div>
            `).join('');

            convertBtn.classList.remove('hidden');
        }

        function removeFile(index) {
            uploadedFiles.splice(index, 1);
            updateFileList();
        }

        function hideResult() {
            document.getElementById('result').classList.add('hidden');
            document.getElementById('progress').classList.add('hidden');
            document.getElementById('progress-text').classList.add('hidden');
        }

        document.getElementById('file-input').addEventListener('change', function(e) {
            const files = Array.from(e.target.files);
            if (tools[currentTool].multiple) {
                uploadedFiles = uploadedFiles.concat(files);
            } else {
                uploadedFiles = files.slice(0, 1);
            }
            updateFileList();
        });

        const uploadArea = document.getElementById('upload-area');
        uploadArea.addEventListener('dragover', function(e) {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });

        uploadArea.addEventListener('dragleave', function(e) {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
        });

        uploadArea.addEventListener('drop', function(e) {
            e.preventDefault();
            uploadArea.classList.remove('dragover');

            const files = Array.from(e.dataTransfer.files);
            if (tools[currentTool].multiple) {
                uploadedFiles = uploadedFiles.concat(files);
            } else {
                uploadedFiles = files.slice(0, 1);
            }
            updateFileList();
        });

        async function convertFiles() {
            if (uploadedFiles.length === 0) return;

            const formData = new FormData();
            uploadedFiles.forEach(file => {
                formData.append('files', file);
            });
            formData.append('tool', currentTool);

            document.getElementById('progress').classList.remove('hidden');
            document.getElementById('progress-text').classList.remove('hidden');
            const btn = document.getElementById('convert-btn');
            btn.disabled = true;
            btn.innerText = 'Processing...';
            
            // Fake progress animation
            const bar = document.getElementById('progress-bar');
            bar.style.width = '30%';
            setTimeout(() => { bar.style.width = '70%'; }, 500);

            try {
                const response = await fetch('/convert', {
                    method: 'POST',
                    body: formData
                });

                if (response.ok) {
                    bar.style.width = '100%';
                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = response.headers.get('Content-Disposition')?.split('filename=')[1] || 'converted_file.zip';
                    document.body.appendChild(a);
                    a.click();
                    window.URL.revokeObjectURL(url);
                    document.body.removeChild(a);

                    document.getElementById('result').innerHTML = '<div class="alert alert-success"><strong>Success!</strong> File processed and downloaded successfully.</div>';
                    document.getElementById('result').classList.remove('hidden');
                } else {
                    throw new Error('Conversion error');
                }
            } catch (error) {
                document.getElementById('result').innerHTML = '<div class="alert alert-error"><strong>Error!</strong> An error occurred while processing. Please try again.</div>';
                document.getElementById('result').classList.remove('hidden');
            } finally {
                setTimeout(() => {
                    document.getElementById('progress').classList.add('hidden');
                    document.getElementById('progress-text').classList.add('hidden');
                    bar.style.width = '0%';
                }, 1000);
                btn.disabled = false;
                btn.innerText = 'Process Files';
            }
        }
    </script>
</body>
</html>
"""


@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)


def excel_to_pdf(file, temp_dir):
    xlsx_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(xlsx_path)

    pdf_path = os.path.join(temp_dir, "excel_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    try:
        workbook = openpyxl.load_workbook(xlsx_path)
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            c.setFont("Helvetica", 10)
            c.drawString(50, y_position, f"--- Planilha: {sheet_name} ---")
            y_position -= 20

            for row_idx, row in enumerate(sheet.iter_rows()):
                row_data = [
                    str(cell.value) if cell.value is not None else "" for cell in row
                ]
                line_text = " | ".join(row_data)

                # Simples quebra de linha para caber na página
                max_line_width = int(
                    (width - 100) / 6
                )  # Estimativa de caracteres por linha
                if len(line_text) > max_line_width:
                    # Implementação mais robusta de quebra de linha seria necessária
                    line_text = line_text[:max_line_width] + "..."

                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                    c.setFont("Helvetica", 10)  # Reset font after new page

                c.drawString(50, y_position, line_text)
                y_position -= 15  # Espaçamento menor para linhas de planilha

            y_position -= 30  # Espaçamento entre planilhas
            if (
                y_position < 50 and sheet_name != workbook.sheetnames[-1]
            ):  # Only show new page if not last sheet
                c.showPage()
                y_position = height - 50

    except Exception as e:
        # Handle potential errors with Excel files
        c.drawString(50, y_position - 20, f"Erro ao ler planilha: {e}")
        print(f"Erro ao ler planilha Excel: {e}")

    c.save()
    return [pdf_path]


def txt_to_pdf(file, temp_dir):
    txt_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(txt_path)

    pdf_path = os.path.join(temp_dir, "text_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    c.setFont("Helvetica", 12)

    try:
        with open(txt_path, "r", encoding="utf-8") as f:
            for line in f:
                # Simples quebra de linha para caber na página
                text_line = line.strip()
                max_width_px = width - 100  # Margens de 50px de cada lado

                # Estimar a largura do texto para quebrar linhas
                # ReportLab não tem quebra automática de texto complexa por default
                # Esta é uma estimativa MUITO simples; para algo robusto, precisaria de TextObject
                approx_char_width_px = 7  # Média para Helvetica 12
                chars_per_line = int(max_width_px / approx_char_width_px)

                if len(text_line) > chars_per_line:
                    # Quebra simples da linha
                    chunks = [
                        text_line[i : i + chars_per_line]
                        for i in range(0, len(text_line), chars_per_line)
                    ]
                else:
                    chunks = [text_line]

                for chunk in chunks:
                    if y_position < 50:  # Margem inferior
                        c.showPage()
                        y_position = height - 50
                        c.setFont("Helvetica", 12)  # Reset font after new page

                    c.drawString(50, y_position, chunk)
                    y_position -= 15  # Espaçamento entre linhas

    except Exception as e:
        c.drawString(50, y_position - 20, f"Erro ao ler arquivo de texto: {e}")
        print(f"Erro ao ler arquivo de texto: {e}")

    c.save()
    return [pdf_path]


@app.route("/convert", methods=["POST"])
def convert():
    if "files" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    files = request.files.getlist("files")
    tool = request.form.get("tool")

    if not files or files[0].filename == "":
        return jsonify({"error": "No file selected"}), 400

    # Validação de extensão dos arquivos enviados
    for f in files:
        if not allowed_file(f.filename):
            return jsonify({"error": f"Extensão não permitida: {f.filename}"}), 400

    # Criar diretório temporário
    temp_dir = tempfile.mkdtemp()
    response = None
    try:
        if tool == "pdf-to-images":
            output_files = pdf_to_images(files[0], temp_dir)
        elif tool == "images-to-pdf":
            output_files = images_to_pdf(files, temp_dir)
        elif tool == "merge-pdf":
            output_files = merge_pdfs(files, temp_dir)
        elif tool == "split-pdf":
            output_files = split_pdf(files[0], temp_dir)
        elif tool == "compress-pdf":
            output_files = compress_pdf(files[0], temp_dir)
        elif tool == "pdf-to-pdfa":
            output_files = pdf_to_pdfa(files, temp_dir)
        elif tool == "word-to-pdf":
            output_files = word_to_pdf(files, temp_dir)
        elif tool == "excel-to-pdf":
            output_files = excel_to_pdf(files[0], temp_dir)
        elif tool == "txt-to-pdf":
            output_files = txt_to_pdf(files[0], temp_dir)
        elif tool == "pdf-to-word":
            output_files = pdf_to_word(files[0], temp_dir)
        elif tool == "ocr-pdf":
            output_files = ocr_pdf(files[0], temp_dir)
        else:
            return jsonify({"error": "Ferramenta não suportada"}), 400

        response = build_response(output_files, temp_dir)
        return response
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        # Diretório temporário limpo após preparar resposta (BytesIO) evitando remoção antecipada
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)


def pdf_to_images(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_files = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x resolution
        img_path = os.path.join(temp_dir, f"page_{page_num + 1}.png")
        pix.save(img_path)
        output_files.append(img_path)

    doc.close()
    return output_files


def images_to_pdf(files, temp_dir):
    images = []
    for file in files:
        img_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(img_path)
        img = Image.open(img_path)
        if img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)

    pdf_path = os.path.join(temp_dir, "images_to_pdf.pdf")
    images[0].save(pdf_path, save_all=True, append_images=images[1:])

    return [pdf_path]


def merge_pdfs(files, temp_dir):
    merged_doc = fitz.open()

    for file in files:
        pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(pdf_path)
        doc = fitz.open(pdf_path)
        merged_doc.insert_pdf(doc)
        doc.close()

    output_path = os.path.join(temp_dir, "merged.pdf")
    merged_doc.save(output_path)
    merged_doc.close()

    return [output_path]


def split_pdf(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_files = []

    for page_num in range(len(doc)):
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
        output_path = os.path.join(temp_dir, f"page_{page_num + 1}.pdf")
        new_doc.save(output_path)
        new_doc.close()
        output_files.append(output_path)

    doc.close()
    return output_files


def compress_pdf(file, temp_dir):
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    doc = fitz.open(pdf_path)
    output_path = os.path.join(temp_dir, "compressed.pdf")
    doc.save(output_path, garbage=4, deflate=True, clean=True)
    doc.close()

    return [output_path]


def pdf_to_pdfa(files, temp_dir):
    """Converte um ou mais PDFs para PDF/A-1b usando Ghostscript."""
    if not isinstance(files, list):
        files = [files]

    output_files = []

    for file in files:
        input_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(input_path)

        base_name, _ = os.path.splitext(os.path.basename(input_path))
        output_path = os.path.join(temp_dir, f"{base_name}_pdfa.pdf")

        gs_args = [
            "gs",
            "-dPDFA=1",
            "-dBATCH",
            "-dNOPAUSE",
            "-dNOOUTERSAVE",
            "-dUseCIEColor",
            "-sProcessColorModel=DeviceRGB",
            "-sDEVICE=pdfwrite",
            "-sColorConversionStrategy=UseDeviceIndependentColor",
            "-dPDFACompatibilityPolicy=1",
            f"-sOutputFile={output_path}",
            input_path,
        ]
        gs_args = [
            arg.encode("utf-8") if isinstance(arg, str) else arg for arg in gs_args
        ]

        try:
            ghostscript.Ghostscript(*gs_args)
        except Exception as e:
            raise RuntimeError(
                f"Erro ao converter {file.filename} para PDF/A: {e}"
            ) from e

        output_files.append(output_path)

    return output_files


def word_to_pdf(files, temp_dir):
    """
    Converte um ou múltiplos arquivos DOCX para PDF
    Se houver múltiplos arquivos, mescla todos em um único PDF
    """
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    # Criar PDF de saída
    pdf_path = os.path.join(temp_dir, "word_to_pdf.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    y_position = height - 50

    # Se for apenas um arquivo (compatibilidade)
    if not isinstance(files, list):
        files = [files]

    # Processar cada arquivo DOCX
    for file_idx, file in enumerate(files):
        docx_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(docx_path)

        # Lê o documento Word
        doc = Document(docx_path)

        # Adicionar separador visual (exceto no primeiro documento)
        if file_idx > 0:
            # Quebra de página
            c.showPage()
            y_position = height - 50

            # Adicionar cabeçalho com nome do arquivo
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y_position, f"{'=' * 60}")
            y_position -= 20
            c.drawString(50, y_position, f"Documento: {file.filename}")
            y_position -= 20
            c.drawString(50, y_position, f"{'=' * 60}")
            y_position -= 30
            c.setFont("Helvetica", 11)

        # Processar parágrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Quebra texto longo em múltiplas linhas
                text = paragraph.text
                max_width = width - 100

                # Estimativa simples de largura de texto
                approx_char_width = 6
                chars_per_line = int(max_width / approx_char_width)

                words = text.split()
                lines = []
                current_line = []

                for word in words:
                    if len(" ".join(current_line + [word])) <= chars_per_line:
                        current_line.append(word)
                    else:
                        if current_line:
                            lines.append(" ".join(current_line))
                            current_line = [word]
                        else:
                            lines.append(word)

                if current_line:
                    lines.append(" ".join(current_line))

                for line in lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50

                    c.drawString(50, y_position, line)
                    y_position -= 20

        # Processar tabelas (se houver)
        for table in doc.tables:
            # Adicionar espaçamento antes da tabela
            y_position -= 10

            if y_position < 100:
                c.showPage()
                y_position = height - 50

            # Desenhar linhas da tabela
            c.setFont("Helvetica", 9)
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])

                # Quebrar texto da linha se necessário
                if len(row_text) > 100:
                    row_text = row_text[:97] + "..."

                if y_position < 50:
                    c.showPage()
                    y_position = height - 50

                c.drawString(50, y_position, row_text)
                y_position -= 15

            # Espaçamento após tabela
            y_position -= 10
            c.setFont("Helvetica", 11)

    c.save()
    return [pdf_path]


def pdf_to_word(file, temp_dir):
    """
    Convert PDF to Word (.docx) format.
    """
    pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(pdf_path)

    docx_filename = os.path.splitext(secure_filename(file.filename))[0] + ".docx"
    docx_path = os.path.join(temp_dir, docx_filename)

    cv = None
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
    except ValueError as e:
        raise RuntimeError(f"Erro no arquivo PDF: {e}") from e
    except ConversionException as e:
        raise RuntimeError(f"Erro interno na conversão: {e}") from e
    except Exception as e:
        raise RuntimeError(f"Erro ao converter {file.filename} para Word: {e}") from e
    finally:
        if cv:
            cv.close()

    return [docx_path]


def ocr_pdf(file, temp_dir):
    """
    Extrai texto de um PDF ou imagem usando Tesseract OCR.
    Retorna um arquivo TXT com o texto extraído.
    """
    filename = secure_filename(file.filename)
    input_path = os.path.join(temp_dir, filename)
    file.save(input_path)

    ext = filename.rsplit(".", 1)[1].lower()
    extracted_text = []

    if ext == "pdf":
        with fitz.open(input_path) as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x resolução
                img_path = os.path.join(temp_dir, f"ocr_page_{page_num + 1}.png")
                pix.save(img_path)

                with Image.open(img_path) as img:
                    text = pytesseract.image_to_string(img, lang="por+eng")
                extracted_text.append(f"--- Página {page_num + 1} ---\n{text}")
    elif ext in ("jpg", "jpeg", "png"):
        # Aplicar OCR diretamente na imagem
        with Image.open(input_path) as img:
            text = pytesseract.image_to_string(img, lang="por+eng")
        extracted_text.append(text)
    else:
        raise RuntimeError(f"Formato não suportado para OCR: {ext}")

    # Salvar texto extraído em arquivo TXT
    base_name = os.path.splitext(filename)[0]
    txt_path = os.path.join(temp_dir, f"{base_name}_ocr.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(extracted_text))

    return [txt_path]


def build_response(output_files, temp_dir):
    """Monta resposta enviando arquivos como attachment sem risco de remoção prematura do diretório temporário."""
    if len(output_files) == 1:
        file_path = output_files[0]
        filename = os.path.basename(file_path)
        with open(file_path, "rb") as f:
            data = f.read()
        return send_file(io.BytesIO(data), as_attachment=True, download_name=filename)
    else:
        zip_path = os.path.join(temp_dir, "converted_files.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for file_path in output_files:
                zipf.write(file_path, os.path.basename(file_path))
        with open(zip_path, "rb") as f:
            data = f.read()
        return send_file(
            io.BytesIO(data), as_attachment=True, download_name="converted_files.zip"
        )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
